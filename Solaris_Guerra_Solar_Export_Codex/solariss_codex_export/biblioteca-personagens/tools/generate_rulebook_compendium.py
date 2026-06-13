#!/usr/bin/env python3
"""Generate a browser-ready compendium from the five final Solaris rulebooks."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


BOOK_META = {
    "book1": ("Livro 1", "Livro Básico do Jogador"),
    "book2": ("Livro 2", "Guia do Mestre"),
    "book3": ("Livro 3", "Bestiário"),
    "book4": ("Livro 4", "Cenários e História"),
    "book5": ("Livro 5", "Itens, Equipamentos e Habilidades"),
}

HEADING_PREFIX = re.compile(
    r"^(?P<number>\d+(?:\.\d+){0,4})[.)]?\s+(?P<title>\S.*)$"
)


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def slug(value: str) -> str:
    normalized = unicodedata.normalize("NFD", clean(value))
    normalized = "".join(
        character
        for character in normalized
        if unicodedata.category(character) != "Mn"
    )
    return re.sub(r"[^a-z0-9]+", "-", normalized.lower()).strip("-") or "secao"


def iter_document_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def paragraph_has_numbering(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    return bool(properties is not None and properties.numPr is not None)


def heading_info(paragraph: Paragraph) -> tuple[int, str, str] | None:
    text = clean(paragraph.text)
    if not text:
        return None
    style = clean(paragraph.style.name if paragraph.style else "")
    normalized_style = slug(style)
    style_match = re.search(r"(\d+)", style)
    styled_heading = (
        normalized_style.startswith("heading")
        or normalized_style.startswith("titulo")
        or normalized_style.startswith("title")
        or normalized_style.startswith("capitulo")
    )
    number_match = HEADING_PREFIX.match(text)
    if styled_heading:
        level = int(style_match.group(1)) if style_match else 1
        number = number_match.group("number") if number_match else ""
        title = number_match.group("title") if number_match else text
        return max(1, min(level, 6)), number, clean(title).rstrip(".")
    if (
        number_match
        and "." in number_match.group("number")
        and len(text) <= 180
        and not paragraph_has_numbering(paragraph)
        and normalized_style not in {"list-paragraph", "paragrafo-de-lista"}
    ):
        number = number_match.group("number")
        level = max(1, min(number.count(".") + 1, 6))
        return level, number, clean(number_match.group("title")).rstrip(".")
    if len(text) <= 140 and re.match(r"^(CAP[IÍ]TULO|LIVRO|AP[EÊ]NDICE)\b", text, re.I):
        return 1, "", text.rstrip(".")
    return None


def paragraph_block(paragraph: Paragraph) -> dict | None:
    text = clean(paragraph.text)
    if not text:
        return None
    style = clean(paragraph.style.name if paragraph.style else "")
    numbered = paragraph_has_numbering(paragraph)
    normalized_style = slug(style)
    kind = "paragraph"
    if numbered or "list" in normalized_style or "lista" in normalized_style:
        kind = "list"
    elif any(token in normalized_style for token in ("citacao", "quote")):
        kind = "quote"
    elif any(token in normalized_style for token in ("nota", "note", "aviso")):
        kind = "note"
    return {"type": kind, "text": text, "style": style}


def table_block(table: Table) -> dict | None:
    rows = [
        [clean(cell.text) for cell in row.cells]
        for row in table.rows
    ]
    rows = [row for row in rows if any(row)]
    if not rows:
        return None
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    return {
        "type": "table",
        "headers": normalized_rows[0],
        "rows": normalized_rows[1:],
    }


def section_summary(blocks: list[dict]) -> str:
    for block in blocks:
        if block["type"] in {"paragraph", "note", "quote"}:
            text = clean(block["text"])
            if len(text) >= 24:
                return text[:360] + ("..." if len(text) > 360 else "")
    for block in blocks:
        if block["type"] == "table":
            headers = [value for value in block["headers"] if value]
            if headers:
                return f"Tabela oficial: {', '.join(headers[:6])}."
    return "Seção oficial do livro de regras Solaris."


def source_label(short_label: str, number: str, breadcrumb: list[str]) -> str:
    if number:
        return f"{short_label}, {number}"
    if breadcrumb:
        return f"{short_label}, {breadcrumb[-1]}"
    return short_label


def build_sections(book_id: str, path: Path) -> tuple[list[dict], dict]:
    short_label, long_label = BOOK_META[book_id]
    document = Document(path)
    sections: list[dict] = []
    current: dict | None = None
    hierarchy: list[dict] = []
    preface_blocks: list[dict] = []

    def close_current() -> None:
        nonlocal current
        if current is None:
            return
        if not current["contentBlocks"]:
            current = None
            return
        current["summary"] = section_summary(current["contentBlocks"])
        sections.append(current)
        current = None

    def start_section(level: int, number: str, title: str) -> None:
        nonlocal current, hierarchy
        close_current()
        hierarchy = [item for item in hierarchy if item["level"] < level]
        hierarchy.append({"level": level, "title": title})
        breadcrumb = [item["title"] for item in hierarchy]
        source = source_label(short_label, number, breadcrumb)
        current = {
            "id": f"{book_id}-{number.replace('.', '-') if number else slug('-'.join(breadcrumb))}-{slug(title)}",
            "category": "rulebook-section",
            "bookId": book_id,
            "bookLabel": short_label,
            "bookTitle": long_label,
            "number": number,
            "title": title,
            "name": " · ".join(filter(None, (short_label, number, title))),
            "level": level,
            "breadcrumb": breadcrumb,
            "tags": [short_label, long_label, "texto oficial"],
            "contentBlocks": [],
            "source": source,
            "schemaVersion": 1,
        }

    for block in iter_document_blocks(document):
        if isinstance(block, Paragraph):
            heading = heading_info(block)
            if heading:
                start_section(*heading)
                continue
            rendered = paragraph_block(block)
        else:
            rendered = table_block(block)
        if not rendered:
            continue
        if current is None:
            preface_blocks.append(rendered)
        else:
            current["contentBlocks"].append(rendered)

    close_current()
    if preface_blocks:
        title = "Apresentação"
        source = short_label
        sections.insert(
            0,
            {
                "id": f"{book_id}-apresentacao",
                "category": "rulebook-section",
                "bookId": book_id,
                "bookLabel": short_label,
                "bookTitle": long_label,
                "number": "",
                "title": title,
                "name": f"{short_label} · {title}",
                "level": 1,
                "breadcrumb": [title],
                "summary": section_summary(preface_blocks),
                "tags": [short_label, long_label, "texto oficial"],
                "contentBlocks": preface_blocks,
                "source": source,
                "schemaVersion": 1,
            },
        )

    stats = {
        "id": book_id,
        "label": short_label,
        "title": long_label,
        "file": path.name,
        "sections": len(sections),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "characters": sum(
            len(block.get("text", ""))
            + sum(len(value) for value in block.get("headers", []))
            + sum(len(value) for row in block.get("rows", []) for value in row)
            for section in sections
            for block in section["contentBlocks"]
        ),
    }
    return sections, stats


def write_output(output: Path, payload: dict) -> None:
    serialized = json.dumps(payload, ensure_ascii=False, indent=2)
    output.write_text(
        "/* Gerado automaticamente a partir dos cinco livros finais de Solaris. */\n"
        f"globalThis.SOLARIS_RULEBOOK_COMPENDIUM = {serialized};\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    for book_id in BOOK_META:
        parser.add_argument(f"--{book_id}", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    sections: list[dict] = []
    sources: list[dict] = []
    for book_id in BOOK_META:
        book_sections, stats = build_sections(book_id, getattr(args, book_id))
        sections.extend(book_sections)
        sources.append(stats)

    id_counts: dict[str, int] = {}
    for section in sections:
        base_id = section["id"]
        id_counts[base_id] = id_counts.get(base_id, 0) + 1
        if id_counts[base_id] > 1:
            section["id"] = f"{base_id}-{id_counts[base_id]}"

    payload = {
        "schemaVersion": 1,
        "sources": sources,
        "sections": sections,
    }
    write_output(args.output, payload)
    print(
        json.dumps(
            {
                "output": str(args.output),
                "sections": len(sections),
                "sources": sources,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()

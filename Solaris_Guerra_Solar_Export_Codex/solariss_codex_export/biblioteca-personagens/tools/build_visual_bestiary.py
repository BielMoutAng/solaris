#!/usr/bin/env python3
"""Build the visual Guerra Solar bestiary from the official Book 3 content."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DATA_PREFIX = "globalThis.SOLARIS_OFFICIAL_BOOKS = "
INK = "0B1620"
DEEP = "071018"
CYAN = "18C9D6"
CYAN_DARK = "0C6D78"
CYAN_PALE = "E8F7F8"
AMBER = "D6A33A"
AMBER_PALE = "FFF4D6"
RED = "C94747"
MUTED = "526774"
LINE = "A9CDD1"
WHITE = "FFFFFF"


def slug(value: str) -> str:
    normalized = unicodedata.normalize("NFD", value)
    ascii_value = "".join(char for char in normalized if unicodedata.category(char) != "Mn")
    return re.sub(r"[^a-z0-9]+", "-", ascii_value.lower()).strip("-")


def read_payload(path: Path) -> dict:
    source = path.read_text(encoding="utf-8")
    start = source.index(DATA_PREFIX) + len(DATA_PREFIX)
    return json.loads(source[start : source.rfind(";")])


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: int = 6) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ARQUIVO XENOBIOLÓGICO  //  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def add_page_border(section, color: str = CYAN_DARK) -> None:
    properties = section._sectPr
    borders = properties.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        properties.append(borders)
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "16")
        border.set(qn("w:color"), color)
        borders.append(border)


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True
    add_page_border(section)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, color, before, after in (
        ("Title", 31, WHITE, 0, 8),
        ("Subtitle", 12, CYAN, 0, 8),
        ("Heading 1", 22, INK, 16, 8),
        ("Heading 2", 15, CYAN_DARK, 12, 5),
        ("Heading 3", 13, INK, 8, 5),
        ("Heading 4", 10.5, CYAN_DARK, 6, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Ficha Campo" not in [style.name for style in document.styles]:
        style = document.styles.add_style("Ficha Campo", 1)
    else:
        style = document.styles["Ficha Campo"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("GUERRA SOLAR  //  LIVRO 3  //  BESTIÁRIO")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    add_page_number(section.footer.paragraphs[0])


def add_label_bar(document: Document, text: str, fill: str = DEEP, color: str = CYAN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [7.15])
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_border(cell, color, 8)
    set_cell_margins(cell, top=95, bottom=95, start=130, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)


def add_cover(document: Document, cover_path: Path) -> None:
    add_label_bar(document, "Solaris // Guerra Solar", DEEP, AMBER)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    run = title.add_run("LIVRO 3")
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("BESTIÁRIO DE GUERRA SOLAR")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(cover_path), height=Inches(7.25))
    edition = document.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.paragraph_format.space_before = Pt(4)
    run = edition.add_run("EDIÇÃO VISUAL OFICIAL  //  ARQUIVO XENOBIOLÓGICO")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(AMBER)
    document.add_page_break()


def add_chapter_divider(document: Document, number: str, title: str, subtitle: str) -> None:
    document.add_page_break()
    add_label_bar(document, f"Capítulo {number}", DEEP, CYAN)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(90)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title.upper())
    run.font.name = "Arial"
    run.font.size = Pt(29)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    run.font.color.rgb = RGBColor.from_string(CYAN)
    paragraph = document.add_paragraph(subtitle)
    paragraph.paragraph_format.space_after = Pt(16)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(callout, [6.7])
    cell = callout.cell(0, 0)
    shade(cell, CYAN_PALE)
    set_cell_border(cell, CYAN_DARK, 7)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PROTOCOLO DE CAMPO // Leia a ameaça, observe seus sinais e escolha como ela transforma a cena.")
    r.font.name = "Arial"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("Sumário operacional", level=1)
    entries = [
        ("1", "Como usar o bestiário", "Leitura das fichas, coleta, sinais e decisões do Mestre."),
        ("2", "Catálogo de ameaças", "Fauna, flora, máquinas, humanoides, corrupção e Cosmos."),
        ("3", "Chefes e ameaças lendárias", "Oito encontros de campanha com fases, arena e consequências."),
        ("4", "Variantes e templates", "Ferramentas para adaptar fichas sem reescrever o monstro."),
        ("5", "Ferramentas do Mestre", "Encontros, sinais, materiais, funções e descobertas."),
        ("6", "Fichas em branco", "Modelos completos, rápidos e de chefe."),
        ("7", "Glossário", "Terminologia oficial do bestiário."),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.55, 2.55, 4.0]
    set_table_width(table, widths)
    for cell, text in zip(table.rows[0].cells, ("ARQ.", "SEÇÃO", "FUNÇÃO")):
        shade(cell, DEEP)
        set_cell_border(cell, CYAN_DARK, 7)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(CYAN)
    set_repeat_table_header(table.rows[0])
    for number, title, description in entries:
        cells = table.add_row().cells
        values = (number, title, description)
        for index, (cell, text) in enumerate(zip(cells, values)):
            shade(cell, CYAN_PALE if int(number) % 2 else WHITE)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.bold = index < 2
            r.font.color.rgb = RGBColor.from_string(CYAN_DARK if index < 2 else INK)
    document.add_paragraph(
        "Esta edição reúne 54 fichas únicas. A ficha revisada da Nyxaracne Matriarca substitui a versão anterior, "
        "e Tyrakth Cristalino é o nome consolidado da variante cristalina."
    )


def add_family_divider(document: Document, title: str, first: bool = False) -> None:
    if not first:
        document.add_page_break()
    descriptions = {
        "Fauna predatória": "Predadores territoriais, caçadores de emboscada e variantes biológicas de Tarantus.",
        "Flora e organismos de zona": "Organismos fixos ou semimóveis que transformam o ambiente em ameaça.",
        "Máquinas e sentinelas": "Drones, guardiões e plataformas antigas movidas por protocolos incompletos.",
        "Humanoides hostis": "Adversários organizados que usam tática, equipamento, medo e objetivos próprios.",
        "Ameaças cósmicas": "Entidades alteradas por sinais, vazio, memória e forças além da matéria comum.",
        "Humanoides corrompidos": "Povos jogáveis deformados pela corrupção, ainda reconhecíveis sob a ameaça.",
    }
    add_label_bar(document, "Classificação de campo", DEEP, CYAN)
    heading = document.add_heading(title, level=2)
    heading.paragraph_format.space_before = Pt(100)
    heading.paragraph_format.space_after = Pt(12)
    heading.runs[0].font.size = Pt(25)
    paragraph = document.add_paragraph(descriptions.get(title, "Grupo de ameaças catalogadas."))
    paragraph.paragraph_format.space_after = Pt(16)
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.7])
    cell = table.cell(0, 0)
    shade(cell, CYAN_PALE)
    set_cell_border(cell, CYAN_DARK, 7)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("LEITURA TÁTICA // Observe habitat, comportamento, fraquezas e recursos antes de iniciar o confronto.")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    document.add_page_break()


def paragraph_texts(document: Document, heading: str) -> list[str]:
    start = None
    start_level = 99
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == heading:
            start = index
            match = re.search(r"(\d+)", paragraph.style.name if paragraph.style else "")
            start_level = int(match.group(1)) if match else 99
            break
    if start is None:
        return []
    texts = []
    for paragraph in document.paragraphs[start + 1 :]:
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading"):
            match = re.search(r"(\d+)", style)
            level = int(match.group(1)) if match else 99
            if level <= start_level:
                break
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if text:
            texts.append(text)
    return texts


def add_source_section(document: Document, source: Document, source_heading: str, title: str | None = None) -> None:
    texts = paragraph_texts(source, source_heading)
    if not texts:
        return
    document.add_heading(title or re.sub(r"^\d+(?:\.\d+)*\.\s*", "", source_heading), level=2)
    for text in texts:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if len(text) < 70 and text.endswith(":"):
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(CYAN_DARK)
        else:
            p.add_run(text)


def normalize_monsters(payload: dict, source: Document) -> tuple[list[dict], list[dict]]:
    monsters = deepcopy(payload["bestiary"])
    revised = next((m for m in monsters if "ficha revisada" in m["name"].lower()), None)
    if revised:
        revised_index = monsters.index(revised)
        original_index = next(
            index
            for index, monster in enumerate(monsters)
            if monster["name"] == "Nyxaracne Matriarca"
        )
        revised["name"] = "Nyxaracne Matriarca"
        revised["id"] = monsters[original_index]["id"]
        revised["source"] = "Livro 3, ficha consolidada de 2.29 e 4.16"
        monsters[original_index] = revised
        monsters = [monster for index, monster in enumerate(monsters) if index != revised_index]

    bosses = [m for m in monsters if m["id"].startswith("livro3-3-")]
    standard = [m for m in monsters if not m["id"].startswith("livro3-3-")]
    for monster, table in zip(bosses, source.tables[5:13]):
        values = {
            row.cells[0].text.strip(): row.cells[1].text.strip()
            for row in table.rows[1:]
            if len(row.cells) >= 2
        }
        monster["tier"] = values.get("Tier", monster.get("tier", ""))
        monster["type"] = values.get("Tipo", monster.get("type", ""))
        monster["role"] = values.get("Papel", monster.get("role", ""))
        monster["size"] = values.get("Tamanho", monster.get("size", ""))
        monster["habitat"] = values.get("Habitat ou arena", monster.get("habitat", ""))
        monster["movement"] = values.get("Movimento", monster.get("movement", ""))
        monster["attributes"] = values.get("Atributos importantes", monster.get("attributes", ""))
        monster["pv"] = first_number(values.get("PV", ""))
        monster["ca"] = first_number(values.get("CA", ""))
        monster["needsCoreStats"] = False
        core = [
            f"Tier: {monster['tier']}.",
            f"Tipo: {monster['type']}.",
            f"Papel: {monster['role']}.",
            f"Tamanho: {monster['size']}.",
            f"Habitat: {monster['habitat']}.",
            f"Objetivo: {values.get('Objetivo', '')}.",
            f"PV: {values.get('PV', '')}.",
            f"CA: {values.get('CA', '')}.",
            f"Movimento: {monster['movement']}.",
            f"Atributos importantes: {monster['attributes']}.",
        ]
        monster["details"] = [{"label": "Dados centrais", "items": core}, *monster.get("details", [])]
    return standard, bosses


def first_number(value: str) -> int | None:
    match = re.search(r"\d+", value or "")
    return int(match.group()) if match else None


def family_name(monster: dict) -> str:
    name = slug(monster["name"])
    if any(key in name for key in ("drone", "guardiao", "sentinela", "juggernautt")):
        return "Máquinas e sentinelas"
    if any(key in name for key in ("saqueador", "cultista-do-sinal")):
        return "Humanoides hostis"
    if "corrompido" in name:
        return "Humanoides corrompidos"
    if any(key in name for key in ("espreitador", "arauto-da-noite")):
        return "Ameaças cósmicas"
    if any(key in name for key in ("fungo", "silvari", "teia-sombra")):
        return "Flora e organismos de zona"
    return "Fauna predatória"


def ordered_standard(monsters: list[dict]) -> list[dict]:
    order = {
        "Fauna predatória": 0,
        "Flora e organismos de zona": 1,
        "Máquinas e sentinelas": 2,
        "Humanoides hostis": 3,
        "Ameaças cósmicas": 4,
        "Humanoides corrompidos": 5,
    }
    indexed = {monster["id"]: index for index, monster in enumerate(monsters)}
    return sorted(monsters, key=lambda monster: (order[family_name(monster)], indexed[monster["id"]]))


def image_path(assets: Path, monster: dict) -> Path:
    return assets / f"{slug(monster['name'])}.jpg"


def add_stat_table(document: Document, monster: dict) -> None:
    labels = ("TIER", "TIPO", "PAPEL", "PV", "CA", "MOVIMENTO")
    values = (
        monster.get("tier") or "—",
        monster.get("type") or "—",
        monster.get("role") or "—",
        str(monster.get("pv") or "—"),
        str(monster.get("ca") or "—"),
        monster.get("movement") or "—",
    )
    table = document.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [0.72, 1.6, 1.45, 0.55, 0.55, 1.45])
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        shade(cell, DEEP)
        set_cell_border(cell, CYAN_DARK, 6)
        set_cell_margins(cell, top=55, bottom=55, start=45, end=45)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(7)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(CYAN)
        value_cell = table.cell(1, index)
        shade(value_cell, CYAN_PALE)
        set_cell_border(value_cell, LINE, 5)
        set_cell_margins(value_cell, top=65, bottom=65, start=45, end=45)
        p = value_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(values[index])
        r.font.name = "Arial"
        r.font.size = Pt(7.5 if index in (1, 2, 5) else 9)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)


def is_core_line(text: str) -> bool:
    return bool(
        re.match(
            r"^(Tier|Tipo|Papel|Tamanho|Habitat(?: ou arena)?|Comportamento|Objetivo|PV|CA|Movimento|Atributos(?: importantes)?)\s*:",
            text,
            re.I,
        )
    )


def add_monster_details(document: Document, monster: dict) -> None:
    seen = set()
    for group in monster.get("details", []):
        items = []
        for item in group.get("items", []):
            text = re.sub(r"\s+", " ", item).strip()
            key = slug(text)
            if not text or is_core_line(text) or key in seen:
                continue
            seen.add(key)
            items.append(text)
        if not items:
            continue
        label = group.get("label", "Informações de campo")
        if label not in ("Resumo", "Dados centrais"):
            p = document.add_paragraph(style="Ficha Campo")
            p.add_run(label)
        for item in items:
            if item.endswith(":") or (
                len(item) < 70
                and any(
                    item.lower().startswith(prefix)
                    for prefix in (
                        "ataque",
                        "habilidade",
                        "resist",
                        "fraque",
                        "sentido",
                        "moral",
                        "recurso",
                        "uso em campanha",
                        "ação",
                        "reação",
                        "fase",
                        "condição",
                        "solução",
                        "consequência",
                    )
                )
            ):
                p = document.add_paragraph(style="Ficha Campo")
                p.add_run(item.rstrip(":"))
            else:
                p = document.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.12)
                p.paragraph_format.first_line_indent = Inches(-0.12)
                p.paragraph_format.space_after = Pt(2.5)
                p.add_run("◆ ").font.color.rgb = RGBColor.from_string(CYAN_DARK)
                p.add_run(item)


def add_monster_entry(
    document: Document,
    monster: dict,
    number: str,
    assets: Path,
    boss: bool = False,
    page_break_before: bool = True,
) -> None:
    heading = document.add_heading(f"{number}. {monster['name']}", level=3 if not boss else 2)
    heading.paragraph_format.page_break_before = page_break_before
    heading.paragraph_format.space_after = Pt(4)

    art = image_path(assets, monster)
    if art.exists():
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_after = Pt(4)
        from PIL import Image

        with Image.open(art) as image:
            ratio = image.width / image.height
        if ratio < 0.9:
            picture.add_run().add_picture(str(art), height=Inches(8.55))
        elif ratio > 1.15:
            picture.add_run().add_picture(str(art), width=Inches(7.0))
        else:
            picture.add_run().add_picture(str(art), height=Inches(7.35))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)
    run = caption.add_run(
        f"REGISTRO VISUAL  //  {monster.get('tier') or 'TIER NÃO CLASSIFICADO'}  //  "
        f"{monster.get('type') or 'AMEAÇA NÃO CLASSIFICADA'}"
    )
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(CYAN_DARK)
    document.add_page_break()

    add_label_bar(document, f"Ficha de campo // {monster['name']}", DEEP, AMBER if boss else CYAN)
    add_stat_table(document, monster)

    core_fields = (
        ("Tier", monster.get("tier")),
        ("Tipo", monster.get("type")),
        ("Papel", monster.get("role")),
        ("Tamanho", monster.get("size")),
        ("Habitat", monster.get("habitat")),
        ("PV", monster.get("pv")),
        ("CA", monster.get("ca")),
        ("Movimento", monster.get("movement")),
        ("Atributos importantes", monster.get("attributes")),
    )
    for label, value in core_fields:
        if value not in (None, ""):
            p = document.add_paragraph(style="Ficha Campo")
            p.add_run(f"{label}: {value}")
    add_monster_details(document, monster)
    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(6)
    source.paragraph_format.keep_together = True
    run = source.add_run(f"FONTE OFICIAL // {monster.get('source', 'Livro 3')}")
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(AMBER)


def copy_source_table(document: Document, source_table, title: str) -> None:
    document.add_heading(title, level=2)
    columns = len(source_table.columns)
    table = document.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    available = 7.05
    widths = [available / columns] * columns
    if columns == 2:
        widths = [0.72, 6.33]
    elif columns == 3:
        widths = [0.72, 2.2, 4.13]
    set_table_width(table, widths)
    for index, source_cell in enumerate(source_table.rows[0].cells):
        cell = table.rows[0].cells[index]
        shade(cell, DEEP)
        set_cell_border(cell, CYAN_DARK, 7)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(source_cell.text.strip())
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(CYAN)
    set_repeat_table_header(table.rows[0])
    for row_index, source_row in enumerate(source_table.rows[1:], start=1):
        cells = table.add_row().cells
        for index, source_cell in enumerate(source_row.cells):
            cell = cells[index]
            shade(cell, CYAN_PALE if row_index % 2 else WHITE)
            set_cell_border(cell)
            set_cell_margins(cell, top=70, bottom=70)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(source_cell.text.strip())
            r.font.name = "Arial"
            r.font.size = Pt(8.3)


def copy_blank_form(document: Document, source_table, title: str) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=0, cols=len(source_table.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [7.05 / len(source_table.columns)] * len(source_table.columns)
    set_table_width(table, widths)
    for source_row in source_table.rows:
        cells = table.add_row().cells
        for index, source_cell in enumerate(source_row.cells):
            cell = cells[index]
            set_cell_border(cell, CYAN_DARK, 5)
            set_cell_margins(cell, top=125, bottom=125, start=90, end=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(source_cell.text.strip())
            r.font.name = "Arial"
            r.font.size = Pt(8)
            r.font.bold = source_row is source_table.rows[0]
            r.font.color.rgb = RGBColor.from_string(CYAN_DARK if r.font.bold else INK)


def add_glossary(document: Document, source: Document) -> None:
    start = next(
        index
        for index, paragraph in enumerate(source.paragraphs)
        if paragraph.text.strip() == "Glossário — Bestiário"
    )
    for paragraph in source.paragraphs[start + 1 :]:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading"):
            document.add_heading(text, level=2)
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2.5)
        if ":" in text and len(text.split(":", 1)[0]) < 45:
            term, definition = text.split(":", 1)
            r = p.add_run(f"{term}:")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(CYAN_DARK)
            p.add_run(definition)
        else:
            p.add_run(text)


def build_document(source_path: Path, data_path: Path, assets: Path, output: Path) -> None:
    source = Document(source_path)
    payload = read_payload(data_path)
    standard, bosses = normalize_monsters(payload, source)
    standard = ordered_standard(standard)

    document = Document()
    set_document_defaults(document)
    document.core_properties.title = "Livro 3 — Bestiário de Guerra Solar"
    document.core_properties.subject = "Bestiário oficial e arquivo xenobiológico de Solaris"
    document.core_properties.author = "Gabriel — Guerra Solar"
    document.core_properties.keywords = "Guerra Solar, Solaris, bestiário, RPG, monstros"

    add_cover(document, assets / "book3-cover.jpg")
    add_toc(document)

    add_chapter_divider(
        document,
        "1",
        "Como usar o bestiário",
        "Procedimentos para ler fichas, apresentar sinais, coletar materiais e transformar monstros em escolhas.",
    )
    document.add_heading("1. Como usar o bestiário", level=1)
    add_source_section(document, source, "2.1. Como usar este bestiário", "1.1. Leitura das fichas")
    add_source_section(document, source, "2.2. Regras gerais de coleta", "1.2. Coleta e análise")
    add_source_section(document, source, "1.16. Variantes de monstros", "1.3. Variantes")
    add_source_section(document, source, "1.17. Monstros como pistas", "1.4. Monstros como pistas")
    add_source_section(document, source, "1.18. Monstros como escolha moral", "1.5. Escolha moral")
    add_source_section(document, source, "1.19. Monstros recorrentes", "1.6. Ameaças recorrentes")
    add_source_section(document, source, "1.20. Regra de ouro dos monstros", "1.7. Regra de ouro")

    add_chapter_divider(
        document,
        "2",
        "Catálogo de ameaças",
        "Fichas únicas organizadas por natureza, função narrativa e origem da ameaça.",
    )
    current_family = None
    first_family = True
    for index, monster in enumerate(standard, start=1):
        family = family_name(monster)
        if family != current_family:
            current_family = family
            add_family_divider(document, current_family, first=first_family)
            first_family = False
            add_monster_entry(document, monster, f"2.{index}", assets, page_break_before=False)
        else:
            add_monster_entry(document, monster, f"2.{index}", assets)

    add_chapter_divider(
        document,
        "3",
        "Chefes e ameaças lendárias",
        "Entidades, máquinas e predadores concebidos como encontros centrais de campanha.",
    )
    for index, monster in enumerate(bosses, start=1):
        add_monster_entry(
            document,
            monster,
            f"3.{index}",
            assets,
            boss=True,
            page_break_before=index != 1,
        )

    add_chapter_divider(
        document,
        "4",
        "Variantes e templates",
        "Camadas de modificação para criar versões únicas sem perder a leitura da ficha original.",
    )
    document.add_heading("4. Variantes e templates", level=1)
    for source_heading, title in (
        ("4.1. Função deste capítulo", "4.1. Função editorial"),
        ("4.2. Padronização de atributos nas fichas", "4.2. Atributos padronizados"),
        ("4.28. Template: Corrompido por Falaris", "4.3. Corrompido por Falaris"),
        ("4.29. Template: Alfa", "4.4. Alfa"),
        ("4.30. Template: Enxame", "4.5. Enxame"),
        ("4.31. Template: Máquina Antiga", "4.6. Máquina Antiga"),
        ("4.32. Template: Predador Noturno", "4.7. Predador Noturno"),
        ("4.33. Template: Guardião Territorial", "4.8. Guardião Territorial"),
        ("4.35. Criaturas por função", "4.9. Criaturas por função"),
        ("4.36. Perícias para descobrir fraquezas", "4.10. Descobrir fraquezas"),
        ("4.37. Regra de ouro do Bestiário Expandido", "4.11. Regra de ouro"),
    ):
        add_source_section(document, source, source_heading, title)

    add_chapter_divider(
        document,
        "5",
        "Ferramentas do Mestre",
        "Tabelas de encontro, sinais, materiais e variações rápidas para preparar cenas.",
    )
    document.add_heading("5. Ferramentas do Mestre", level=1)
    copy_source_table(document, source.tables[0], "5.1. Encontros por ambiente")
    copy_source_table(document, source.tables[1], "5.2. Sinais antes da ameaça")
    copy_source_table(document, source.tables[2], "5.3. Materiais coletáveis")
    copy_source_table(document, source.tables[3], "5.4. Variantes rápidas")
    copy_source_table(document, source.tables[13], "5.5. Sinais de chefe")
    add_source_section(document, source, "2.46. Criaturas opcionais por ambiente", "5.6. Criaturas por ambiente")
    add_source_section(document, source, "2.51. Usando monstros como consequência", "5.7. Monstros como consequência")
    add_source_section(document, source, "2.52. Usando monstros como pistas", "5.8. Monstros como pistas")
    add_source_section(document, source, "2.53. Regra de ouro do bestiário", "5.9. Regra de ouro")

    add_chapter_divider(
        document,
        "6",
        "Fichas em branco",
        "Modelos prontos para registrar criaturas comuns, ameaças rápidas e chefes.",
    )
    document.add_heading("6. Fichas em branco", level=1)
    copy_blank_form(document, source.tables[14], "6.1. Ficha completa de criatura")
    document.add_page_break()
    copy_blank_form(document, source.tables[15], "6.2. Ficha rápida de criatura")
    document.add_page_break()
    copy_blank_form(document, source.tables[16], "6.3. Ficha de chefe")

    add_chapter_divider(
        document,
        "7",
        "Glossário",
        "Vocabulário oficial para consulta rápida durante preparação e jogo.",
    )
    document.add_heading("7. Glossário do bestiário", level=1)
    add_glossary(document, source)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(
        json.dumps(
            {
                "output": str(output),
                "standard": len(standard),
                "bosses": len(bosses),
                "total": len(standard) + len(bosses),
            },
            ensure_ascii=False,
        )
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--data", type=Path, required=True)
    parser.add_argument("--assets", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_document(args.source, args.data, args.assets, args.output)


if __name__ == "__main__":
    main()

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guia_GitHub_e_Executaveis_Solaris.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6874"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "A87513"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering_definition(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_list(doc, items, num_id):
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        apply_numbering(paragraph, num_id)
        set_run_font(paragraph.add_run(item))


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "101820")
    p_pr.append(shading)
    for index, line in enumerate(text.strip().splitlines()):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9, color="E8F2F7")
        if index < len(text.strip().splitlines()) - 1:
            run.add_break()
    return paragraph


def add_callout(doc, title, text, color=BLUE, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title + " ")
    set_run_font(title_run, size=10.5, color=color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660])
    for index, (label, detail) in enumerate(rows):
        set_cell_shading(table.cell(index, 0), LIGHT_BLUE)
        left = table.cell(index, 0).paragraphs[0]
        right = table.cell(index, 1).paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        right.paragraph_format.space_after = Pt(0)
        set_run_font(left.add_run(label), size=10, color=INK, bold=True)
        set_run_font(right.add_run(detail), size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("SOLARIS | GUIA DE PUBLICAÇÃO"), size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("Página "), size=9, color=MUTED)
    add_page_field(fp)


def build_document():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    bullet_id = add_numbering_definition(doc, bullet=True)

    doc.core_properties.title = "Guia prático de GitHub e executáveis"
    doc.core_properties.subject = "Publicação de projetos e distribuição do Solaris"
    doc.core_properties.author = "Projeto Solaris"
    doc.core_properties.keywords = "GitHub, Git, Electron, executável, release, Solaris"

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    set_run_font(kicker.add_run("MANUAL PRÁTICO"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("GitHub e Executáveis"), size=29, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run_font(subtitle.add_run("Como publicar projetos, criar versões para download e distribuir aplicativos"), size=14, color=DARK_BLUE)
    purpose = doc.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose.paragraph_format.space_after = Pt(90)
    set_run_font(purpose.add_run("Exemplo aplicado: Biblioteca de Personagens Solaris v0.3.0"), size=11, color=MUTED, italic=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("Junho de 2026"), size=11, color=INK, bold=True)
    doc.add_page_break()

    doc.add_heading("Como usar este guia", level=1)
    doc.add_paragraph(
        "Este manual foi escrito para você repetir o processo sozinho. Ele começa pelo fluxo mais simples, explica o que cada comando faz e termina com a automação que transforma uma tag de versão em uma página de download no GitHub."
    )
    add_label_table(doc, [
        ("Parte 1", "Publicar e atualizar um projeto no GitHub."),
        ("Parte 2", "Transformar um projeto em aplicativo executável."),
        ("Parte 3", "Criar Releases automáticas com instaladores para download."),
        ("Exemplo", "O projeto Solaris usa Electron e gera instalador NSIS e versão portátil para Windows."),
    ])
    add_callout(
        doc,
        "Ideia central:",
        "Git guarda o histórico no seu computador. GitHub recebe uma cópia desse histórico na internet. Uma Release aponta para uma versão específica e oferece arquivos prontos para baixar.",
    )

    doc.add_heading("1. Preparação inicial", level=1)
    doc.add_heading("1.1 Contas e programas", level=2)
    add_list(doc, [
        "Crie uma conta em https://github.com.",
        "Instale o Git para Windows em https://git-scm.com/download/win.",
        "Instale o Node.js LTS quando o projeto usar JavaScript, Electron, React ou ferramentas npm.",
        "Opcional: instale o GitHub Desktop se preferir botões em vez de comandos.",
        "Opcional: instale o GitHub CLI para criar Releases e Pull Requests pelo terminal.",
    ], bullet_id)
    add_callout(doc, "Teste rápido:", "Abra o PowerShell e execute os comandos abaixo. Se ambos mostrarem versões, a preparação básica está pronta.")
    add_code(doc, "git --version\nnode --version\nnpm --version")

    doc.add_heading("1.2 Configurar seu nome no Git", level=2)
    add_code(doc, 'git config --global user.name "Seu Nome"\ngit config --global user.email "seu-email@exemplo.com"')
    doc.add_paragraph("Use de preferência o mesmo e-mail cadastrado no GitHub. Essa configuração identifica quem criou cada commit.")

    doc.add_heading("2. Publicar um projeto novo no GitHub", level=1)
    doc.add_heading("2.1 Criar o repositório no site", level=2)
    add_list(doc, [
        "Entre no GitHub e clique em New repository.",
        "Escolha um nome curto e sem espaços.",
        "Defina se será público ou privado.",
        "Ao enviar uma pasta que já existe, não marque a criação automática de README, .gitignore ou licença.",
        "Clique em Create repository e mantenha a página aberta para copiar a URL.",
    ], add_numbering_definition(doc, bullet=False))

    doc.add_heading("2.2 Preparar a pasta local", level=2)
    doc.add_paragraph("No PowerShell, entre na pasta do projeto. Coloque o caminho entre aspas quando houver espaços.")
    add_code(doc, 'cd "C:\\caminho\\do\\seu-projeto"\ngit init\ngit status')
    doc.add_paragraph("Depois crie ou revise o arquivo .gitignore. Ele impede que arquivos temporários, senhas, dependências e instaladores locais sejam enviados.")
    add_code(doc, "node_modules/\ndist/\n.env\n.vscode/\n**/__pycache__/")
    add_callout(doc, "Nunca publique:", "senhas, tokens, chaves privadas, arquivos .env, credenciais de banco, certificados ou dados pessoais de jogadores.", color=RED, fill="FDECEC")

    doc.add_heading("2.3 Criar o primeiro commit", level=2)
    add_code(doc, 'git add .\ngit status\ngit commit -m "Primeira versão do projeto"')
    add_label_table(doc, [
        ("git add .", "Prepara os arquivos alterados para o próximo commit."),
        ("git status", "Mostra o que será incluído e ajuda a evitar arquivos indesejados."),
        ("git commit", "Cria um ponto do histórico com uma mensagem."),
    ])

    doc.add_heading("2.4 Conectar e enviar ao GitHub", level=2)
    add_code(doc, "git branch -M main\ngit remote add origin https://github.com/SEU-USUARIO/SEU-REPOSITORIO.git\ngit push -u origin main")
    doc.add_paragraph("Na primeira autenticação, o Git pode abrir o navegador. Autorize a conta correta. Depois do primeiro push, normalmente basta usar git push.")

    doc.add_heading("3. Atualizar um projeto já publicado", level=1)
    doc.add_paragraph("O fluxo diário recomendado é curto e previsível:")
    add_code(doc, 'git status\ngit add caminho/do/arquivo\ngit commit -m "Descreve claramente a alteração"\ngit push')
    add_callout(doc, "Boa prática:", "Prefira adicionar arquivos específicos. Use git add . somente depois de revisar git status.")

    doc.add_heading("3.1 Trabalhar em uma branch", level=2)
    add_code(doc, 'git switch -c codex/nova-funcionalidade\n# faça e teste as alterações\ngit add .\ngit commit -m "Adiciona nova funcionalidade"\ngit push -u origin codex/nova-funcionalidade')
    doc.add_paragraph("No GitHub, abra um Pull Request para comparar a branch com main, revisar e mesclar as mudanças.")

    doc.add_heading("3.2 Baixar mudanças feitas por outra pessoa", level=2)
    add_code(doc, "git switch main\ngit pull origin main")
    add_callout(doc, "Antes de usar git pull:", "salve ou faça commit das suas alterações. Isso reduz conflitos e evita misturar trabalhos incompletos.")

    doc.add_heading("4. Versões e tags", level=1)
    doc.add_paragraph("Use versionamento semântico no formato MAIOR.MENOR.CORREÇÃO.")
    add_label_table(doc, [
        ("1.0.0", "Primeira versão estável."),
        ("1.1.0", "Nova funcionalidade compatível."),
        ("1.1.1", "Correção sem mudança grande de comportamento."),
        ("2.0.0", "Mudança incompatível ou grande reformulação."),
    ])
    doc.add_heading("4.1 Criar uma tag manualmente", level=2)
    add_code(doc, 'git tag -a v0.3.0 -m "Solaris v0.3.0"\ngit push origin v0.3.0')
    doc.add_paragraph("A tag fixa um nome em um commit. No Solaris, enviar uma tag que começa com v também inicia a criação automática dos executáveis.")

    doc.add_heading("5. O que é um executável", level=1)
    doc.add_paragraph(
        "Um projeto web normalmente precisa de navegador e servidor. Um aplicativo desktop empacota a interface e o runtime necessários. A tecnologia usada define a ferramenta de empacotamento."
    )
    add_label_table(doc, [
        ("Electron", "Aplicações HTML, CSS e JavaScript. Gera .exe com electron-builder."),
        ("Python", "Scripts e aplicações Python. PyInstaller é uma opção comum."),
        ("Java", "Pode gerar JAR e instaladores com jpackage."),
        ("Aplicação web", "Pode ser publicada como site ou PWA, sem necessariamente gerar .exe."),
    ])

    doc.add_heading("6. Transformar um projeto Electron em executável", level=1)
    doc.add_paragraph("A Biblioteca Solaris já está configurada como aplicativo Electron. Os passos abaixo também servem como modelo para outros projetos JavaScript.")
    doc.add_heading("6.1 Estrutura mínima", level=2)
    add_list(doc, [
        "package.json com main apontando para o arquivo principal do Electron.",
        "Arquivo electron-main.cjs responsável por criar a janela.",
        "Arquivos da interface, como index.html, styles.css e app.js.",
        "electron e electron-builder instalados como dependências de desenvolvimento.",
        "Configuração build definindo nome, identificador, arquivos incluídos e alvos.",
    ], bullet_id)

    doc.add_heading("6.2 Instalar dependências e testar", level=2)
    add_code(doc, "npm ci\nnpm test\nnpm start")
    doc.add_paragraph("npm ci instala exatamente as versões do package-lock.json. npm test executa os testes. npm start abre o aplicativo em modo desktop para conferência.")

    doc.add_heading("6.3 Gerar o instalador", level=2)
    add_code(doc, "npm run dist")
    doc.add_paragraph("O Solaris usa dois alvos para Windows:")
    add_list(doc, [
        "NSIS: instalador tradicional, permite escolher pasta e cria atalhos.",
        "Portable: executável único que pode ser aberto sem instalação.",
    ], bullet_id)
    add_callout(doc, "Saída:", "Os arquivos aparecem na pasta dist. Essa pasta é resultado de build e não deve ser enviada diretamente ao histórico do Git.")

    doc.add_heading("6.4 Exemplo de scripts do package.json", level=2)
    add_code(doc, '"scripts": {\n  "start": "electron .",\n  "test": "node --test tests/solaris-domain-architecture.test.mjs",\n  "pack": "electron-builder --dir",\n  "dist": "electron-builder --win nsis portable"\n}')

    doc.add_heading("7. Criar executável para um projeto Python", level=1)
    doc.add_paragraph("Quando o projeto é Python, uma alternativa simples é o PyInstaller.")
    add_code(doc, "python -m venv .venv\n.\\.venv\\Scripts\\Activate.ps1\npip install -r requirements.txt\npip install pyinstaller\npyinstaller --onefile --name MeuPrograma app.py")
    add_callout(doc, "Interfaces gráficas:", "Use --windowed quando o programa não precisar exibir um terminal. Teste o arquivo em outro computador antes de publicar.")

    doc.add_heading("8. Releases automáticas no GitHub", level=1)
    doc.add_paragraph(
        "GitHub Actions executa tarefas em máquinas do GitHub. O Solaris possui um workflow que inicia quando uma tag v* é enviada."
    )
    add_list(doc, [
        "Baixa o código da tag.",
        "Prepara o Node.js.",
        "Instala dependências com npm ci.",
        "Executa os testes.",
        "Gera instalador e versão portátil no Windows.",
        "Cria a GitHub Release e anexa os arquivos .exe.",
    ], add_numbering_definition(doc, bullet=False))
    doc.add_heading("8.1 Fluxo completo de uma nova versão", level=2)
    add_code(doc, 'npm version 0.3.1 --no-git-tag-version\ngit add .\ngit commit -m "Prepara Solaris v0.3.1"\ngit push\ngit tag -a v0.3.1 -m "Solaris v0.3.1"\ngit push origin v0.3.1')
    doc.add_paragraph("Depois, abra a aba Actions no GitHub. Quando o trabalho terminar, a nova versão aparecerá na seção Releases.")

    doc.add_heading("8.2 Onde seus amigos baixam", level=2)
    add_list(doc, [
        "Abra o repositório no GitHub.",
        "Na coluna direita, clique em Releases.",
        "Abra a versão mais recente.",
        "Em Assets, escolha o instalador ou a versão portátil.",
        "No Windows, confirme o aviso de segurança somente se o arquivo veio do seu repositório oficial.",
    ], add_numbering_definition(doc, bullet=False))

    doc.add_heading("9. Como publicar manualmente uma Release", level=1)
    add_list(doc, [
        "No repositório, clique em Releases e depois Draft a new release.",
        "Escolha uma tag existente ou crie uma nova.",
        "Escreva um título como Solaris v0.3.0.",
        "Descreva novidades, correções e eventuais limitações.",
        "Arraste os arquivos do diretório dist para a área de anexos.",
        "Clique em Publish release.",
    ], add_numbering_definition(doc, bullet=False))
    add_callout(doc, "Quando usar:", "A publicação manual é útil se o workflow automático falhar ou se você quiser anexar materiais extras, como PDF, ficha de exemplo ou changelog.")

    doc.add_heading("10. O que foi configurado no Solaris v0.3.0", level=1)
    add_list(doc, [
        "Exclusão individual e em massa com confirmação e sem reembolso de Luzentis.",
        "Proteção contra exclusão de arma ou armadura equipada.",
        "Tratamento especial para armazenadores com conteúdo.",
        "Instalação, desinstalação e exclusão de chips modificadores.",
        "Remoção de magias, chip de profissão e mods instalados.",
        "Ficha jogável de monstro com PV, CA, ataques, dano, cura, condições e notas.",
        "Paginação automática de listas com limite de 20 cards.",
        "Criação de personagem aleatório de nível 1.",
        "Migração segura de itens antigos para localização não atribuída.",
        "Workflow de Release para Windows e testes automatizados.",
    ], bullet_id)

    doc.add_heading("11. Checklist antes de publicar", level=1)
    add_list(doc, [
        "O projeto abre sem erros.",
        "Os testes passam.",
        "O número da versão foi atualizado.",
        "git status mostra apenas arquivos esperados.",
        "Nenhuma senha ou chave foi incluída.",
        "O executável foi testado localmente.",
        "A tag corresponde à versão do package.json.",
        "A GitHub Action terminou em verde.",
        "Os arquivos aparecem em Assets na Release.",
        "O link de download foi testado.",
    ], bullet_id)

    doc.add_heading("12. Solução de problemas", level=1)
    add_label_table(doc, [
        ("remote origin already exists", "Use git remote -v para conferir. Para trocar: git remote set-url origin URL."),
        ("rejected / non-fast-forward", "Execute git pull, resolva conflitos, teste e envie novamente."),
        ("nothing to commit", "Não há alterações novas preparadas. Confira git status."),
        ("npm ci falha", "Confirme se package.json e package-lock.json estão sincronizados."),
        ("electron-builder falha", "Apague somente a pasta dist, reinstale dependências e repita npm run dist."),
        ("Release sem arquivos", "Abra Actions, leia o passo de build e confira o caminho configurado em files."),
        ("Windows bloqueia o app", "Aplicativos sem assinatura podem gerar alerta. Distribua somente pelo repositório oficial."),
    ])

    doc.add_heading("13. Comandos de referência rápida", level=1)
    add_code(doc, "# Ver estado\ngit status\n\n# Preparar e salvar\ngit add .\ngit commit -m \"Mensagem clara\"\n\n# Enviar\ngit push\n\n# Atualizar\ngit pull\n\n# Criar versão\nnpm version 0.3.1 --no-git-tag-version\ngit tag -a v0.3.1 -m \"Versão v0.3.1\"\ngit push origin v0.3.1\n\n# Gerar executável Electron\nnpm ci\nnpm test\nnpm run dist")

    doc.add_heading("14. Regra de ouro", level=1)
    add_callout(
        doc,
        "Publique com calma:",
        "revise git status, execute os testes, gere o instalador localmente e só então crie a tag. Esse pequeno ritual evita a maior parte dos problemas de distribuição.",
        color=DARK_BLUE,
        fill=LIGHT_BLUE,
    )
    doc.add_paragraph(
        "Com esse fluxo, seus projetos deixam de ser apenas pastas no seu computador: passam a ter histórico, colaboração, versões identificáveis e downloads reproduzíveis."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
